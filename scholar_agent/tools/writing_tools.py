"""Controlled document tools for paper writing and polishing workflows."""

from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Iterable

from scholar_agent.config import conf
from scholar_agent.core.logging import get_logger
from scholar_agent.tools.base import BaseTool, ToolResult

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:  # pragma: no cover - exercised only when dependency is absent
    Document = None
    Inches = None

_DOC_SUFFIXES = {".docx", ".tex", ".md", ".txt"}
_LATEX_SUFFIXES = {".tex", ".bib", ".cls", ".sty", ".bst"}
_LATEX_ENGINES = {"auto", "latexmk", "pdflatex", "xelatex", "lualatex"}


def _workspace_root() -> Path:
    root = Path(conf.WRITING_WORKSPACE_DIR).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _resolve_document_path(path: str) -> Path:
    root = _workspace_root()
    rel = Path(str(path or "").replace("\\", "/"))
    if rel.is_absolute() or ".." in rel.parts:
        raise ValueError("path must be relative to the writing workspace and must not contain '..'")
    target = (root / rel).resolve()
    if root != target and root not in target.parents:
        raise ValueError("resolved path escapes the writing workspace")
    return target


def _resolve_child_document_path(base_dir: Path, reference: str) -> Path:
    rel = Path(str(reference or "").strip().strip('"').strip("'").replace("\\", "/"))
    if not rel.suffix:
        rel = rel.with_suffix(".tex")
    if rel.is_absolute() or ".." in rel.parts:
        raise ValueError(f"LaTeX include path is outside the writing workspace: {reference}")
    root = _workspace_root()
    target = (base_dir / rel).resolve()
    if root != target and root not in target.parents:
        raise ValueError(f"LaTeX include path escapes the writing workspace: {reference}")
    return target


def _relative_path(path: Path) -> str:
    return str(path.relative_to(_workspace_root())).replace(os.sep, "/")


def _limit_text(text: str, max_chars: int | None) -> str:
    limit = max(1000, min(int(max_chars or 20000), 200000))
    return text[:limit]


def _require_python_docx() -> None:
    if Document is None or Inches is None:
        raise RuntimeError("python-docx is not installed; install dependency 'python-docx' first")


def _iter_docx_blocks(document) -> Iterable[object]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _format_docx_paragraph(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    style_name = paragraph.style.name if paragraph.style is not None else ""
    heading_match = re.fullmatch(r"Heading\s+([1-6])", style_name or "", flags=re.IGNORECASE)
    if heading_match:
        return f"{'#' * int(heading_match.group(1))} {text}"
    if style_name.lower().startswith("title"):
        return f"# {text}"
    return text


def _format_docx_table(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([" ".join(cell.text.split()) for cell in row.cells])
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    markdown_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(cell or " " for cell in row) + " |" for row in markdown_rows)


def _read_docx_text(path: Path, include_tables: bool = True) -> str:
    _require_python_docx()
    document = Document(str(path))
    parts: list[str] = []
    for block in _iter_docx_blocks(document):
        if block.__class__.__name__ == "Paragraph":
            text = _format_docx_paragraph(block)
        else:
            text = _format_docx_table(block) if include_tables else ""
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _configure_docx_page(document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)


def _add_docx_content(document, content: str) -> None:
    lines = str(content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    added = False
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[\.)]\s+(.+)$", stripped)
        if heading:
            document.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 4))
        elif bullet:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
        elif numbered:
            document.add_paragraph(numbered.group(1).strip(), style="List Number")
        else:
            document.add_paragraph(stripped)
        added = True
    if not added:
        document.add_paragraph("")


def _write_docx(path: Path, content: str, title: str = "", author: str = "") -> None:
    _require_python_docx()
    document = Document()
    _configure_docx_page(document)
    if title:
        document.add_heading(str(title).strip(), level=0)
    if author:
        document.add_paragraph(str(author).strip())
    _add_docx_content(document, content)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _latex_include_references(content: str) -> list[str]:
    pattern = re.compile(r"\\(?:input|include|subfile)\s*\{([^}]+)\}")
    return [match.group(1).strip() for match in pattern.finditer(content)]


def _read_latex_text(path: Path, include_inputs: bool = False, max_files: int = 20) -> str:
    visited: set[Path] = set()
    sections: list[str] = []

    def read_one(target: Path) -> None:
        if target in visited or len(visited) >= max_files or not target.exists():
            return
        if target.suffix.lower() not in _LATEX_SUFFIXES:
            return
        visited.add(target)
        text = _read_text(target)
        sections.append(f"% --- {_relative_path(target)} ---\n{text}")
        if include_inputs and target.suffix.lower() == ".tex":
            for reference in _latex_include_references(text):
                try:
                    read_one(_resolve_child_document_path(target.parent, reference))
                except ValueError as exc:
                    sections.append(f"% Skipped unsafe include {reference}: {exc}")

    read_one(path)
    return "\n\n".join(sections)


def _tail_output(text: str, max_chars: int = 12000) -> str:
    text = str(text or "")
    if len(text) <= max_chars:
        return text
    return text[-max_chars:]


def _latex_command(path: Path, engine: str) -> tuple[list[str], str]:
    requested = (engine or "auto").strip().lower()
    if requested not in _LATEX_ENGINES:
        raise ValueError(f"engine must be one of: {', '.join(sorted(_LATEX_ENGINES))}")

    latexmk = shutil.which("latexmk")
    xelatex = shutil.which("xelatex")
    lualatex = shutil.which("lualatex")
    pdflatex = shutil.which("pdflatex")

    if requested == "auto":
        if latexmk:
            if xelatex:
                return [latexmk, "-xelatex", "-interaction=nonstopmode", "-halt-on-error", path.name], "latexmk"
            return [latexmk, "-pdf", "-interaction=nonstopmode", "-halt-on-error", path.name], "latexmk"
        requested = "xelatex" if xelatex else "lualatex" if lualatex else "pdflatex"

    executable = {"latexmk": latexmk, "xelatex": xelatex, "lualatex": lualatex, "pdflatex": pdflatex}.get(requested)
    if not executable:
        raise RuntimeError(f"LaTeX engine is not available on PATH: {requested}")
    if requested == "latexmk":
        return [executable, "-pdf", "-interaction=nonstopmode", "-halt-on-error", path.name], requested
    return [executable, "-interaction=nonstopmode", "-halt-on-error", "-file-line-error", path.name], requested


class WritingListDocumentsTool(BaseTool):
    name = "writing_list_documents"
    description = "List .docx, .tex, .bib, .cls, .sty, .bst, .md, and .txt files under the controlled writing workspace."
    params = {"type": "object", "properties": {}, "additionalProperties": False}

    def execute(self) -> ToolResult:
        root = _workspace_root()
        files = [
            str(path.relative_to(root))
            for path in sorted(root.rglob("*"))
            if path.is_file() and path.suffix.lower() in (_DOC_SUFFIXES | _LATEX_SUFFIXES)
        ]
        return ToolResult("success", "\n".join(files) or "No writing documents found.", data=files)


class WritingReadDocxDocumentTool(BaseTool):
    name = "writing_read_docx_document"
    description = (
        "Read a .docx document from the controlled writing workspace using python-docx. "
        "Returns Markdown-like text preserving headings, paragraphs, bullet/numbered text, and simple tables."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative .docx path under the writing workspace."},
            "max_chars": {
                "type": "integer",
                "description": "Maximum characters to return, default 30000.",
                "default": 30000,
            },
            "include_tables": {
                "type": "boolean",
                "description": "Whether to include simple tables as Markdown tables.",
                "default": True,
            },
        },
        "required": ["path"],
        "additionalProperties": False,
    }

    def execute(self, path: str, max_chars: int = 30000, include_tables: bool = True) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() != ".docx":
                return ToolResult("fail", "DOCX input path must end in .docx.")
            if not target.exists():
                return ToolResult("fail", f"Document not found: {path}")
            text = _read_docx_text(target, include_tables=bool(include_tables))
            return ToolResult(
                "success",
                _limit_text(text, max_chars),
                data={"path": str(target), "chars": len(text), "format": "docx"},
            )
        except Exception as exc:
            return ToolResult("fail", f"DOCX read failed: {exc}")


class WritingReadLatexDocumentTool(BaseTool):
    name = "writing_read_latex_document"
    description = (
        "Read a LaTeX source file from the controlled writing workspace. "
        "Optionally follows \\input, \\include, and \\subfile references within the workspace."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative .tex/.bib/.cls/.sty/.bst path."},
            "max_chars": {
                "type": "integer",
                "description": "Maximum characters to return, default 40000.",
                "default": 40000,
            },
            "include_inputs": {
                "type": "boolean",
                "description": "Whether to recursively include referenced .tex files.",
                "default": False,
            },
        },
        "required": ["path"],
        "additionalProperties": False,
    }

    def execute(self, path: str, max_chars: int = 40000, include_inputs: bool = False) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() not in _LATEX_SUFFIXES:
                return ToolResult("fail", "LaTeX input path must end in .tex, .bib, .cls, .sty, or .bst.")
            if not target.exists():
                return ToolResult("fail", f"Document not found: {path}")
            text = _read_latex_text(target, include_inputs=bool(include_inputs))
            return ToolResult(
                "success",
                _limit_text(text, max_chars),
                data={"path": str(target), "chars": len(text), "format": target.suffix.lower().lstrip(".")},
            )
        except Exception as exc:
            return ToolResult("fail", f"LaTeX read failed: {exc}")


class WritingReadDocumentTool(BaseTool):
    name = "writing_read_document"
    description = (
        "Read a .docx, .tex, .bib, .cls, .sty, .bst, .md, or .txt document from the controlled writing workspace. "
        "For richer format-specific behavior, prefer writing_read_docx_document or writing_read_latex_document."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative path under the writing workspace."},
            "max_chars": {
                "type": "integer",
                "description": "Maximum characters to return, default 20000.",
                "default": 20000,
            },
        },
        "required": ["path"],
        "additionalProperties": False,
    }

    def execute(self, path: str, max_chars: int = 20000) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if not target.exists():
                return ToolResult("fail", f"Document not found: {path}")
            suffix = target.suffix.lower()
            if suffix == ".docx":
                text = _read_docx_text(target)
            elif suffix in _LATEX_SUFFIXES | {".md", ".txt"}:
                text = _read_latex_text(target) if suffix in _LATEX_SUFFIXES else _read_text(target)
            else:
                return ToolResult("fail", "Supported formats are .docx, .tex, .bib, .cls, .sty, .bst, .md, and .txt.")
            return ToolResult("success", _limit_text(text, max_chars), data={"path": str(target), "chars": len(text)})
        except Exception as exc:
            return ToolResult("fail", f"Read failed: {exc}")


class WritingWriteTextDocumentTool(BaseTool):
    name = "writing_write_text_document"
    description = (
        "Write a .tex, .bib, .cls, .sty, .bst, .md, or .txt document to the controlled writing workspace. "
        "For LaTeX source files, prefer writing_write_latex_document."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative output path ending in a supported text extension."},
            "content": {"type": "string", "description": "Full document content to write."},
        },
        "required": ["path", "content"],
        "additionalProperties": False,
    }

    def execute(self, path: str, content: str) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() not in (_LATEX_SUFFIXES | {".md", ".txt"}):
                return ToolResult("fail", "Text output path must end in .tex, .bib, .cls, .sty, .bst, .md, or .txt.")
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(str(content or ""), encoding="utf-8")
            logger.info("[tool] writing_write_text_document path=%s", target)
            return ToolResult("success", f"Wrote {_relative_path(target)}", data={"path": str(target)})
        except Exception as exc:
            return ToolResult("fail", f"Write failed: {exc}")


class WritingWriteDocxDocumentTool(BaseTool):
    name = "writing_write_docx_document"
    description = (
        "Write a .docx document to the controlled writing workspace using python-docx. "
        "Accepts Markdown-like plain text; # headings become Word headings, -/* bullets become list items, and numbered lines become numbered lists."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative output path ending in .docx."},
            "content": {"type": "string", "description": "Markdown-like text content to write."},
            "title": {"type": "string", "description": "Optional title inserted before content.", "default": ""},
            "author": {"type": "string", "description": "Optional author line inserted below title.", "default": ""},
        },
        "required": ["path", "content"],
        "additionalProperties": False,
    }

    def execute(self, path: str, content: str, title: str = "", author: str = "") -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() != ".docx":
                return ToolResult("fail", "DOCX output path must end in .docx.")
            _write_docx(target, content, title=title or "", author=author or "")
            logger.info("[tool] writing_write_docx_document path=%s", target)
            return ToolResult("success", f"Wrote {_relative_path(target)}", data={"path": str(target)})
        except Exception as exc:
            return ToolResult("fail", f"DOCX write failed: {exc}")


class WritingWriteLatexDocumentTool(BaseTool):
    name = "writing_write_latex_document"
    description = (
        "Write a LaTeX-related source file to the controlled writing workspace. "
        "Use for .tex, .bib, .cls, .sty, and .bst files."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative output path ending in .tex, .bib, .cls, .sty, or .bst."},
            "content": {"type": "string", "description": "Full LaTeX source content to write."},
        },
        "required": ["path", "content"],
        "additionalProperties": False,
    }

    def execute(self, path: str, content: str) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() not in _LATEX_SUFFIXES:
                return ToolResult("fail", "LaTeX output path must end in .tex, .bib, .cls, .sty, or .bst.")
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(str(content or ""), encoding="utf-8")
            logger.info("[tool] writing_write_latex_document path=%s", target)
            return ToolResult("success", f"Wrote {_relative_path(target)}", data={"path": str(target)})
        except Exception as exc:
            return ToolResult("fail", f"LaTeX write failed: {exc}")


class WritingCompileLatexDocumentTool(BaseTool):
    name = "writing_compile_latex_document"
    description = (
        "Compile a .tex document inside the controlled writing workspace with latexmk, xelatex, lualatex, or pdflatex. "
        "Returns success only when the compiler exits with code 0 and the expected PDF exists."
    )
    params = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Relative main .tex file path under the writing workspace."},
            "engine": {
                "type": "string",
                "description": "LaTeX engine: auto, latexmk, xelatex, lualatex, or pdflatex.",
                "enum": ["auto", "latexmk", "xelatex", "lualatex", "pdflatex"],
                "default": "auto",
            },
            "timeout_s": {
                "type": "integer",
                "description": "Compile timeout in seconds, default 60.",
                "default": 60,
            },
        },
        "required": ["path"],
        "additionalProperties": False,
    }

    def execute(self, path: str, engine: str = "auto", timeout_s: int = 60) -> ToolResult:
        try:
            target = _resolve_document_path(path)
            if target.suffix.lower() != ".tex":
                return ToolResult("fail", "LaTeX compile input path must end in .tex.")
            if not target.exists():
                return ToolResult("fail", f"Document not found: {path}")

            command, resolved_engine = _latex_command(target, engine)
            timeout = max(5, min(int(timeout_s or 60), 300))
            logger.info("[tool] writing_compile_latex_document path=%s engine=%s", target, resolved_engine)
            completed = subprocess.run(
                command,
                cwd=str(target.parent),
                text=True,
                capture_output=True,
                timeout=timeout,
                check=False,
            )
            output = _tail_output((completed.stdout or "") + "\n" + (completed.stderr or ""))
            pdf_path = target.with_suffix(".pdf")
            if completed.returncode == 0 and pdf_path.exists():
                return ToolResult(
                    "success",
                    f"Compiled {_relative_path(target)} with {resolved_engine}; PDF: {_relative_path(pdf_path)}\n{output}".strip(),
                    data={
                        "path": str(target),
                        "pdf_path": str(pdf_path),
                        "engine": resolved_engine,
                        "returncode": completed.returncode,
                    },
                )
            return ToolResult(
                "fail",
                (
                    f"LaTeX compile failed for {_relative_path(target)} with {resolved_engine}; "
                    f"returncode={completed.returncode}; pdf_exists={pdf_path.exists()}\n{output}"
                ).strip(),
                data={
                    "path": str(target),
                    "pdf_path": str(pdf_path),
                    "engine": resolved_engine,
                    "returncode": completed.returncode,
                    "pdf_exists": pdf_path.exists(),
                },
            )
        except subprocess.TimeoutExpired as exc:
            return ToolResult("fail", f"LaTeX compile timed out after {exc.timeout} seconds.")
        except Exception as exc:
            return ToolResult("fail", f"LaTeX compile failed: {exc}")
